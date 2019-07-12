#! /usr/bin/python
# customInvitations.py - Generates a Word document with custom invitations for guest names in a text file.

import docx

def customInvitations(txtFile, blankDoc):
    nameFile = open(txtFile)
    names = [name.strip('\n') for name in nameFile.readlines()]
    nameFile.close()
    doc = docx.Document(blankDoc)
    for name in names:
        doc.add_paragraph('It would be a pleasure to have the company of', 'New1')
        doc.add_paragraph(name, 'New2')
        doc.add_paragraph('at 11010 Memory Lane of the Evening of', 'New1')
        doc.add_paragraph('April 1st', 'New3')
        doc.add_paragraph('at 7 o\'clock', 'New1')
        doc.paragraphs[-1].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
    doc.save(blankDoc)

customInvitations('guests.txt', 'invite.docx')